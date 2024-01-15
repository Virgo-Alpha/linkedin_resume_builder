import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_hyperlink(run, url, text):
    """
    A function that places a hyperlink within a Run object.

    :param run: The Run object we are adding the hyperlink to.
    :param url: A string containing the required url.
    :param text: The text displayed for the url.
    :return: The hyperlink object.
    """
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), run.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True))

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add formatting to the hyperlink (optional)
    rStyle = docx.oxml.OxmlElement('w:rStyle')
    rStyle.set(docx.oxml.shared.qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)

    # Append the hyperlink to the Run's parent
    run._element.append(hyperlink)

    return hyperlink

# document = Document()
# p = document.add_paragraph()
# add_hyperlink(p, 'http://www.baidu.com', 'baidu')
# document.save('demo_hyperlink.docx')

def add_horizontal_line(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run("_" * 105)
    run.font.size = Pt(8)
    run.bold = True

def set_paragraph_spacing(paragraph, before=None, after=None):
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
