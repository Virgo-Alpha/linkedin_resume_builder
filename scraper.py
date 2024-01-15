from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from bs4 import BeautifulSoup
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from demo_link import add_hyperlink, set_paragraph_spacing, add_horizontal_line

def scrape_contact_info(soup):
    contact_info = soup.find_all('section', {'class': 'pv-contact-info__contact-type'})
    contact_dict = {}

    for contact in contact_info:
        if 'linkedin' in contact.text.strip():
            contact_dict["LinkedIn"] = contact.find('a').text.strip()

        if 'Email' in contact.text.strip():
            contact_dict["Email"] = contact.find('a').text.strip()

        if 'Website' in contact.text.strip():
            for link in contact.find_all('li'):
                if 'Portfolio' in link.find('span').text.strip():
                    contact_dict["Portfolio"] = link.find('a').text.strip()

                elif 'Blog' in link.find('span').text.strip():
                    contact_dict["Blog"] = link.find('a').text.strip()

                elif 'Personal' in link.find('span').text.strip():
                    contact_dict["Personal Page"] = link.find('a').text.strip()

        if 'Phone' in contact.text.strip():
            contact_dict["Phone"] = contact.find('a').text.strip()

        if 'Address' in contact.text.strip():
            contact_dict["Address"] = contact.find('a').text.strip()

    return contact_dict

def scrape_basic_info(soup):
    intro = soup.find('div', {'class': 'mt2 relative'})
    name = intro.find("h1").get_text().strip()
    works_at_loc = intro.find("div", {'class': 'text-body-medium'})
    works_at = works_at_loc.get_text().strip()
    location_loc = intro.find("span", {'class': 'text-body-small inline t-black--light break-words'})
    location = location_loc.get_text().strip()

    return {"Name": name, "Works At": works_at, "Location": location}

def scrape_about(soup):
    about_div = soup.find("div", {"id": "about"})
    about_items = about_div.find_parent("section", {"class": "artdeco-card"}).find_all("span", {"aria-hidden": True})

    paragraph = about_items[1].text.strip()
    sentences = paragraph.split('.')
    about = '.'.join(sentences[:2]) + '.'

    return {"About": about}

def scrape_education(soup):
    education_list = []
    education_div = soup.find("div", {"id": "education"})

    if education_div is not None:
        parent_section = education_div.find_parent("section", {"class": "artdeco-card"})
        education_items = parent_section.find_all("li", class_="artdeco-list__item")

        for item in education_items:
            program_details = item.find("span", class_="t-14 t-normal")
            program_name = program_details.find("span", class_="visually-hidden").text.strip()
            school_details = item.find("div", class_="t-bold")
            school = school_details.find("span", class_="visually-hidden").text.strip()
            duration_detail = item.find("span", class_="t-14 t-normal t-black--light")
            duration = duration_detail.find("span", class_="visually-hidden").text.strip().split("·")[0].strip()

            education_entry = {
                "Program": program_name,
                "School": school,
                "Duration": duration
            }
            education_list.append(education_entry)

    return {"Education": education_list}

def scrape_experience(soup):
    experience_list = []
    experience_div = soup.find("div", {"id": "experience"})

    if experience_div is not None:
        parent_section = experience_div.find_parent("section", {"class": "artdeco-card"})
        experience_items = parent_section.find_all("li", class_="artdeco-list__item")

        for item in experience_items[:3]:
            company_details = item.find("span", class_="t-14 t-normal")
            company_name = company_details.find("span", class_="visually-hidden").text.strip()
            position_detail = item.find("div", class_="t-bold")
            position = position_detail.find("span", class_="visually-hidden").text.strip()
            duration_detail = item.find("span", class_="t-14 t-normal t-black--light")
            duration = duration_detail.find("span", class_="visually-hidden").text.strip().split("·")[0].strip()
            description_details = item.find_next("div", class_="pv-shared-text-with-see-more")
            description = description_details.find("span", class_="visually-hidden").text.strip().split("·")[0].strip()

            experience_entry = {
                "Company": company_name,
                "Position": position,
                "Duration": duration,
                "Description": description
            }
            experience_list.append(experience_entry)

    return {"Experience": experience_list}

def scrape_skills(soup):
    skill_list = []
    skills_div = soup.find("div", {"id": "skills"})

    if skills_div is not None:
        parent_section = skills_div.find_parent("section", {"class": "artdeco-card"})
        skill_items = parent_section.find_all("span", {"aria-hidden": True})

        for item in skill_items[1:]:
            parent_div = item.find_parent("div", {"class": "mr1"})
            if parent_div is not None:
                for skill in item.text.strip().split(' • '):
                    skill_list.append(skill)

    return {"Skills": skill_list}

def scrape_languages(soup):
    languages_div = soup.find("div", {"id": "languages"})

    if languages_div is not None:
        language_list = []
        parent_section = languages_div.find_parent("section", {"class": "artdeco-card"})

        # Find the span elements containing the languages within the parent section
        languages = parent_section.find_all("div", {"class": "t-bold"})

        for language in languages:
            real_language = language.find("span", {"class": "visually-hidden"})
            language_list.append(real_language.get_text())

        return {"Languages": language_list}

def scrape_interests(soup):
    interests_div = soup.find("div", {"id": "volunteer_causes"})

    if interests_div is not None:
        interest_list = []
        parent_section = interests_div.find_parent("section", {"class": "artdeco-card"})

        # Find the span elements containing the languages within the parent section
        Causes_items = parent_section.find_all("span", {"aria-hidden": True})

        for item in Causes_items[1:]:
            for cause in item.text.strip().split(' • '):
                interest_list.append(cause)

    return {"Interests": interest_list}

def create_resume(info):
    document = Document()

    # Add a heading with the person's name
    name = info["Name"]

    heading = document.add_heading(name)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(heading, after=1)

   # Add a section for contact info
    contact_list = []

    if "LinkedIn" in info:
        contact_list.append(("LinkedIn", info['LinkedIn']))

    if "Email" in info:
        contact_list.append(("Email", info['Email']))

    if "Portfolio" in info:
        contact_list.append(("Portfolio", info['Portfolio']))

    if "Blog" in info:
        contact_list.append(("Blog", info['Blog']))

    if "Personal Page" in info:
        contact_list.append(("Personal Page", info['Personal Page']))

    if "Phone" in info:
        contact_list.append(("Phone", info['Phone']))

    if "Address" in info:
        contact_list.append(("Address", info['Address']))

    # Calculate the split point based on the length of the contact_list
    split_point = len(contact_list) // 2 + len(contact_list) % 2

    # Create two lists
    list1 = document.add_paragraph()
    # list1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for name, value in contact_list[:split_point]:
        if name in ["LinkedIn", "Portfolio", "Blog", "Personal Page"]:
            # hyperlink = run.add_hyperlink(value, color=(0, 0, 255), underline=False)
            
            hyperlink = add_hyperlink(list1, f"https://{value}", f"{name}")
            # print(hyperlink)
        else:
            run = list1.add_run(f"{value}")

        list1.add_run(" | ")

    list2 = document.add_paragraph()
    for name, value in contact_list[split_point:]:
        if name in ["LinkedIn", "Portfolio", "Blog", "Personal Page"]:
            # hyperlink = run.add_hyperlink(value, color=(0, 0, 255), underline=False)
            add_hyperlink(list2, f"https://{value}", f"{name}")
        else:
            run = list2.add_run(f"{value}")

        list2.add_run(" | ")

    set_paragraph_spacing(list1, after=0)
    set_paragraph_spacing(list2, after=0)

    # ! # Add a section for works at and location
    # document.add_heading("Professional Details", level=2)
    # document.add_paragraph(f"Works At: {info['Works At']}")
    # document.add_paragraph(f"Location: {info['Location']}")

    # Add the About section
    about = document.add_paragraph(info["About"])
    set_paragraph_spacing(about, after=0)

    # Add the Education section
    education_heading = document.add_paragraph("Education")
    education_heading.runs[0].bold = True
    add_horizontal_line(document)
    set_paragraph_spacing(education_heading, after=0)

    for education in info["Education"]:
    # Add School and Duration in two columns with bold formatting
        education_paragraph = document.add_paragraph(
            f"{education['School']}\t\t\t{education['Duration']}",
            style="BodyText"
        )
        education_paragraph.runs[0].bold = True

        # Add Program in a separate paragraph
        program_paragraph = document.add_paragraph(f"{education['Program']}")
        set_paragraph_spacing(program_paragraph, after=0)
        # document.add_paragraph("\n")

    # Add the Experience section
    experience_heading = document.add_paragraph("Work Experience")
    experience_heading.runs[0].bold = True
    add_horizontal_line(document)
    set_paragraph_spacing(experience_heading, after=0)

    for experience in info["Experience"]:
        # Add Position and Duration in two columns with bold formatting
        experience_paragraph = document.add_paragraph(
            f"{experience['Position']}\t\t\t{experience['Duration']}",
            style="BodyText"
        )
        experience_paragraph.runs[0].bold = True

        # Add Company in bold and no columns
        company_paragraph = document.add_paragraph(f"{experience['Company']}")
        company_paragraph.runs[0].bold = True
        set_paragraph_spacing(company_paragraph, after=0)
        # document.add_paragraph("\n")

        # Add Description
        description_paragraph = document.add_paragraph(f"{experience['Description']}")
        set_paragraph_spacing(description_paragraph, after=0)
        # document.add_paragraph("\n")


    # Add the Skills section
    skills_heading = document.add_paragraph("Skills")
    skills_heading.runs[0].bold = True
    add_horizontal_line(document)
    set_paragraph_spacing(skills_heading, after=0)

    skills_paragraph = document.add_paragraph(", ".join(info["Skills"]))
    set_paragraph_spacing(skills_paragraph, after=0)

    # Add the Languages section
    languages_heading = document.add_paragraph("Languages")
    languages_heading.runs[0].bold = True
    add_horizontal_line(document)
    set_paragraph_spacing(languages_heading, after=0)

    languages_paragraph = document.add_paragraph(", ".join(info["Languages"]))
    set_paragraph_spacing(languages_paragraph, after=0)

    # Add the Interests section
    interests_heading = document.add_paragraph("Interests")
    interests_heading.runs[0].bold = True
    add_horizontal_line(document)
    set_paragraph_spacing(interests_heading, after=0)

    interests_paragraph = document.add_paragraph(", ".join(info["Interests"]))
    set_paragraph_spacing(interests_paragraph, after=0)

    # Save the document
    
    # document.save(f"{name}_resume.docx")

    document.save('resume.docx')

def main(profile_url, password, email):
    # Creating a webdriver instance in headless mode
    chrome_options = Options()
    chrome_options.add_argument("--headless")  # This line makes the browser run in headless mode
    driver = webdriver.Chrome(options=chrome_options)
    # driver = webdriver.Chrome()

    try:
        # TODO: Use LinkedIn API to Sign in instead of logging in
        # Logging into LinkedIn
        driver.get("https://linkedin.com/uas/login")
        time.sleep(5)
        username = driver.find_element(By.ID, "username")
        username.send_keys(email)
        pword = driver.find_element(By.ID, "password")
        pword.send_keys(password)
        driver.find_element(By.XPATH, "//button[@type='submit']").click()
        time.sleep(5)

        # Opening the Profile
        driver.get(profile_url)
        time.sleep(5)

        # Clicking the "Contact info" link
        contact_info_link = driver.find_element(By.ID, 'top-card-text-details-contact-info')
        contact_info_link.click()
        time.sleep(5)

        # Scraping contact info
        src = driver.page_source
        soup = BeautifulSoup(src, 'lxml')

        contact_dict = scrape_contact_info(soup)
        
        # Close the modal
        close_button = driver.find_element(By.CLASS_NAME, 'artdeco-modal__dismiss')

        close_button.click()

        # ! we need to scroll the page to get the complete data
        start = time.time()

        # will be used in the while loop
        initialScroll = 0
        finalScroll = 1000

        while True:
            driver.execute_script(f"window.scrollTo({initialScroll}, {finalScroll})")
            # this command scrolls the window starting from
            # the pixel value stored in the initialScroll 
            # variable to the pixel value stored at the
            # finalScroll variable
            initialScroll = finalScroll
            finalScroll += 1000

            # we will stop the script for 3 seconds so that 
            # the data can load
            time.sleep(3)
            # You can change it as per your needs and internet speed

            end = time.time()

            # We will scroll for 20 seconds.
            # You can change it as per your needs and internet speed
            if round(end - start) > 20:
                break

        basic_info_dict = scrape_basic_info(soup)
        about_dict = scrape_about(soup)
        education_dict = scrape_education(soup)
        experience_dict = scrape_experience(soup)
        skills_dict = scrape_skills(soup)
        languages_dict = scrape_languages(soup)
        interests_dict = scrape_interests(soup)

        # Combine all dictionaries into one
        info = {**basic_info_dict, **contact_dict, **about_dict, **education_dict, **experience_dict, **skills_dict, **languages_dict, **interests_dict}

        # Create the resume
        create_resume(info)

    finally:
        # Close the driver
        driver.quit()
